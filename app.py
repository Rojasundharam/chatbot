import os
from datetime import datetime, timedelta
from typing import List, Dict, Any
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import io
from embedding_utils import EmbeddingUtil
import logging
from anthropic import Anthropic
from dotenv import load_dotenv
import docx

load_dotenv()  # Load environment variables from .env file

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class EnhancedDocument:
    def __init__(self, doc_id: str, content: str, metadata: Dict[str, Any]):
        self.doc_id = doc_id
        self.content = content
        self.metadata = metadata
        self.chunks: List[str] = []
        self.embeddings: List[List[float]] = []
        self.last_updated: datetime = datetime.now()

    def chunk_content(self, chunk_size: int = 1000):
        self.chunks = [self.content[i:i+chunk_size] for i in range(0, len(self.content), chunk_size)]

    def create_embeddings(self, embedding_util: EmbeddingUtil):
        self.embeddings = embedding_util.create_embeddings(self.chunks)

class ChatBot:
    def __init__(self, session_state):
        self.session_state = session_state
        self.drive_service = self._get_drive_service()
        self.folder_id = "1EyR0sfFEBUDGbPn3lBDIP5qcFumItrvQ"  # Your Google Drive folder ID
        self.documents: Dict[str, EnhancedDocument] = {}
        self.embedding_util = EmbeddingUtil()
        
        anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")
        if not anthropic_api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable is not set")
        self.anthropic = Anthropic(api_key=anthropic_api_key)
        
        self.load_or_update_documents()

    def _get_drive_service(self):
        creds_path = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')
        if not creds_path:
            raise ValueError("GOOGLE_APPLICATION_CREDENTIALS environment variable is not set")
        creds = service_account.Credentials.from_service_account_file(
            creds_path,
            scopes=['https://www.googleapis.com/auth/drive.readonly']
        )
        return build('drive', 'v3', credentials=creds)

    def load_or_update_documents(self):
        files = self._list_files_in_folder()
        logging.info(f"Found {len(files)} files in the folder")
        for file in files:
            try:
                if file['id'] not in self.documents or self._needs_update(file):
                    content = self._get_file_content(file['id'], file['name'])
                    logging.info(f"Processing file: {file['name']}, Content length: {len(content)}")
                    doc = EnhancedDocument(file['id'], content, file)
                    doc.chunk_content()
                    doc.create_embeddings(self.embedding_util)
                    self.documents[file['id']] = doc
                    logging.info(f"Successfully processed file: {file['name']}, Chunks: {len(doc.chunks)}, Embeddings: {len(doc.embeddings)}")
                else:
                    logging.info(f"Skipped processing file (already up to date): {file['name']}")
            except Exception as e:
                logging.error(f"Error processing file {file['name']}: {str(e)}")
        logging.info(f"Total documents processed: {len(self.documents)}")

    def _list_files_in_folder(self):
        results = self.drive_service.files().list(
            q=f"'{self.folder_id}' in parents",
            fields="files(id, name, modifiedTime, mimeType)"
        ).execute()
        return results.get('files', [])

    def _get_file_content(self, file_id: str, file_name: str) -> str:
        request = self.drive_service.files().get_media(fileId=file_id)
        file = io.BytesIO()
        downloader = MediaIoBaseDownload(file, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        file.seek(0)
        
        if file_name.lower().endswith('.docx'):
            doc = docx.Document(file)
            return '\n'.join([para.text for para in doc.paragraphs])
        else:
            return file.read().decode('utf-8', errors='ignore')

    def _needs_update(self, file: Dict[str, Any]) -> bool:
        if file['id'] not in self.documents:
            return True
        last_modified = datetime.fromisoformat(file['modifiedTime'].rstrip('Z'))
        return last_modified > self.documents[file['id']].last_updated

    def expand_query(self, query):
        expanded_terms = {
            "admission": ["enrollment", "application", "entry requirements", "how to apply", "admission process"],
            "course": ["program", "curriculum", "syllabus", "subjects", "degree", "certificate"],
            "arts": ["fine arts", "liberal arts", "humanities", "visual arts", "performing arts"],
            "engineering": ["technology", "technical courses", "B.Tech", "M.Tech"],
            "pharmacy": ["pharmaceutical", "B.Pharm", "M.Pharm", "drug studies"],
            "nursing": ["healthcare", "B.Sc Nursing", "M.Sc Nursing", "patient care"],
            "dental": ["dentistry", "BDS", "MDS", "oral health"],
            "allied health": ["physiotherapy", "occupational therapy", "medical laboratory"],
            "fees": ["tuition", "cost", "financial aid", "scholarship"],
            "faculty": ["professors", "lecturers", "teaching staff", "academics"],
            "facilities": ["infrastructure", "labs", "library", "hostel", "campus"],
        }
        
        expanded_query = query.lower()
        for term, expansions in expanded_terms.items():
            if term in expanded_query:
                expanded_query += " " + " ".join(expansions)
        
        expanded_query += " JKKN J.K.K. Nattraja educational institution"
        return expanded_query

    def get_relevant_context(self, query: str, max_chunks: int = 5) -> str:
        query_embedding = self.embedding_util.create_embeddings([query])[0]
        relevant_chunks = []
        for doc in self.documents.values():
            similarities = [self.embedding_util.compute_similarity(query_embedding, chunk_embedding) 
                            for chunk_embedding in doc.embeddings]
            top_indices = sorted(range(len(similarities)), key=lambda i: similarities[i], reverse=True)[:max_chunks]
            doc_chunks = [(doc.chunks[i], similarities[i], doc.metadata['name']) for i in top_indices]
            relevant_chunks.extend(doc_chunks)
        
        relevant_chunks.sort(key=lambda x: x[1], reverse=True)
        selected_chunks = relevant_chunks[:max_chunks]
        
        context = "\n\n".join([chunk[0] for chunk in selected_chunks])
        logging.info(f"Selected chunks from documents: {[chunk[2] for chunk in selected_chunks]}")
        logging.info(f"Total relevant context length: {len(context)}")
        return context

    def process_user_input(self, user_input: str) -> str:
        try:
            logging.info(f"Processing user input: {user_input}")
            
            # Fallback for greetings
            greetings = ["hi", "hello", "hey", "greetings"]
            if user_input.lower() in greetings:
                return "Hello! Welcome to JKKN Assist. How can I help you with information about our educational institutions today?"
            
            expanded_query = self.expand_query(user_input)
            logging.info(f"Expanded query: {expanded_query}")
            context = self.get_relevant_context(expanded_query)
            
            if not context.strip():
                return "I'm sorry, but I couldn't find any specific information related to your query in the JKKN documents. Can you please ask about our courses, admissions, facilities, or any specific JKKN institution?"

            logging.info(f"Context found (first 500 characters): {context[:500]}")

            rag_message = f"""Based ONLY on the following information from JKKN institutional documents, answer the user's question:

Context:
{context}

User Question: {user_input}

Instructions:
1. Use ONLY the information provided in the context above to answer the question.
2. If the context doesn't contain specific information that answers the user's question, say so.
3. Do not include any information or assumptions that are not explicitly stated in the provided context.
4. Begin your response with "According to the JKKN documents:" to emphasize that the information comes directly from the institution's materials.

Answer:
"""

            response_message = self.generate_message([{"role": "user", "content": rag_message}])
            assistant_response = response_message.content[0].text
            logging.info(f"Generated response: {assistant_response[:100]}...")
            return assistant_response
        except Exception as e:
            logging.error(f"Error processing user input: {str(e)}")
            return "I apologize, but I encountered an error while processing your request. Could you please try rephrasing your question about JKKN institutions?"

    def generate_message(self, messages, max_tokens=2048):
        try:
            response = self.anthropic.messages.create(
                model="claude-2.1",  # or whichever model you're using
                max_tokens=max_tokens,
                messages=messages
            )
            return response
        except Exception as e:
            logging.error(f"Error generating message: {str(e)}")
            raise