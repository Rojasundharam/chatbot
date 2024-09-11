import os
from datetime import datetime
from typing import List, Dict, Any
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import io
import logging
from anthropic import Anthropic
from dotenv import load_dotenv
import docx
import numpy as np
from sentence_transformers import SentenceTransformer
from rank_bm25 import BM25Okapi
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

load_dotenv()
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class EnhancedDocument:
    def __init__(self, doc_id: str, content: str, metadata: Dict[str, Any]):
        self.doc_id = doc_id
        self.content = content
        self.metadata = metadata
        self.chunks: List[str] = []
        self.last_updated: datetime = datetime.now()

    def chunk_content(self, chunk_size: int = 1000):
        self.chunks = [self.content[i:i+chunk_size] for i in range(0, len(self.content), chunk_size)]

class AdvancedRetrieval:
    def __init__(self, documents: List[EnhancedDocument]):
        self.documents = documents
        self.chunks = [chunk for doc in documents for chunk in doc.chunks]
        
        # Dense retrieval model
        self.dense_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.dense_embeddings = self.dense_model.encode(self.chunks)
        
        # Sparse retrieval model
        self.bm25 = BM25Okapi(self.chunks)
        
        # TF-IDF for re-ranking
        self.tfidf = TfidfVectorizer()
        self.tfidf_matrix = self.tfidf.fit_transform(self.chunks)

    def hybrid_search(self, query: str, k: int = 10) -> List[str]:
        # Dense retrieval
        query_embedding = self.dense_model.encode([query])[0]
        dense_scores = cosine_similarity([query_embedding], self.dense_embeddings)[0]
        
        # Sparse retrieval
        sparse_scores = np.array(self.bm25.get_scores(query.split()))
        
        # Combine scores (you can adjust the weights)
        combined_scores = 0.7 * dense_scores + 0.3 * sparse_scores
        
        # Get top-k results
        top_k_indices = combined_scores.argsort()[-k:][::-1]
        return [self.chunks[i] for i in top_k_indices]

    def rerank(self, query: str, candidates: List[str], n: int = 5) -> List[str]:
        query_vector = self.tfidf.transform([query])
        candidate_indices = [self.chunks.index(c) for c in candidates]
        candidate_vectors = self.tfidf_matrix[candidate_indices]
        
        similarities = cosine_similarity(query_vector, candidate_vectors)[0]
        reranked_indices = similarities.argsort()[-n:][::-1]
        
        return [candidates[i] for i in reranked_indices]

    def get_relevant_context(self, query: str, max_chunks: int = 5) -> str:
        expanded_query = expand_query(query)
        candidates = self.hybrid_search(expanded_query, k=max_chunks*2)
        reranked_chunks = self.rerank(expanded_query, candidates, n=max_chunks)
        return "\n\n".join(reranked_chunks)

class ChatBot:
    def __init__(self, session_state):
        self.session_state = session_state
        self.drive_service = self._get_drive_service()
        self.folder_id = "1EyR0sfFEBUDGbPn3lBDIP5qcFumItrvQ"  # Your Google Drive folder ID
        self.documents: List[EnhancedDocument] = []
        
        anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")
        if not anthropic_api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable is not set")
        self.anthropic = Anthropic(api_key=anthropic_api_key)
        
        self.load_or_update_documents()
        self.retrieval = AdvancedRetrieval(self.documents)

    def _get_drive_service(self):
        creds_path = os.getenv('GOOGLE_APPLICATION_CREDENTIALS')
        if not creds_path:
            raise ValueError("GOOGLE_APPLICATION_CREDENTIALS environment variable is not set")
        creds = service_account.Credentials.from_service_account_file(
            creds_path,
            scopes=['https://www.googleapis.com/auth/drive.readonly']
        )
        return build('drive', 'v3', credentials=creds)

    def load_or_update_documents(self):
        files = self._list_files_in_folder()
        logging.info(f"Found {len(files)} files in the folder")
        for file in files:
            try:
                if not any(doc.doc_id == file['id'] for doc in self.documents) or self._needs_update(file):
                    content = self._get_file_content(file['id'], file['name'])
                    logging.info(f"Processing file: {file['name']}, Content length: {len(content)}")
                    doc = EnhancedDocument(file['id'], content, file)
                    doc.chunk_content()
                    self.documents.append(doc)
                    logging.info(f"Successfully processed file: {file['name']}, Chunks: {len(doc.chunks)}")
                else:
                    logging.info(f"Skipped processing file (already up to date): {file['name']}")
            except Exception as e:
                logging.error(f"Error processing file {file['name']}: {str(e)}")
        logging.info(f"Total documents processed: {len(self.documents)}")

    def _list_files_in_folder(self):
        results = self.drive_service.files().list(
            q=f"'{self.folder_id}' in parents",
            fields="files(id, name, modifiedTime, mimeType)"
        ).execute()
        return results.get('files', [])

    def _get_file_content(self, file_id: str, file_name: str) -> str:
        request = self.drive_service.files().get_media(fileId=file_id)
        file = io.BytesIO()
        downloader = MediaIoBaseDownload(file, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        file.seek(0)
        
        if file_name.lower().endswith('.docx'):
            doc = docx.Document(file)
            return '\n'.join([para.text for para in doc.paragraphs])
        else:
            return file.read().decode('utf-8', errors='ignore')

    def _needs_update(self, file: Dict[str, Any]) -> bool:
        for doc in self.documents:
            if doc.doc_id == file['id']:
                last_modified = datetime.fromisoformat(file['modifiedTime'].rstrip('Z'))
                return last_modified > doc.last_updated
        return True

    def process_user_input(self, user_input: str) -> str:
        try:
            logging.info(f"Processing user input: {user_input}")
            context = self.retrieval.get_relevant_context(user_input)
            
            if not context.strip():
                logging.warning("No relevant context found")
                return "I'm sorry, but I couldn't find any specific information related to your query in the JKKN documents."

            logging.info(f"Context found (first 500 characters): {context[:500]}")

            rag_message = f"""Based ONLY on the following information from JKKN institutional documents, answer the user's question:

Context:
{context}

User Question: {user_input}

Instructions:
1. Use ONLY the information provided in the context above to answer the question.
2. If the context doesn't contain specific information that answers the user's question, say so.
3. Do not include any information or assumptions that are not explicitly stated in the provided context.
4. Begin your response with "According to the JKKN documents:" to emphasize that the information comes directly from the institution's materials.

Answer:
"""

            response_message = self.generate_message([{"role": "user", "content": rag_message}])
            assistant_response = response_message.content[0].text
            logging.info(f"Generated response: {assistant_response[:100]}...")
            return assistant_response
        except Exception as e:
            logging.error(f"Error processing user input: {str(e)}")
            return "I apologize, but I encountered an error while processing your request. Could you please try rephrasing your question about JKKN institutions?"

    def generate_message(self, messages, max_tokens=2048):
        try:
            response = self.anthropic.messages.create(
                model="claude-2.1",
                max_tokens=max_tokens,
                messages=messages
            )
            return response
        except Exception as e:
            logging.error(f"Error generating message: {str(e)}")
            raise

def expand_query(query: str) -> str:
    expansions = {
        "eng": "engineering",
        "fee": "fees tuition cost",
        "admission": "admissions enrollment application",
        "course": "course program curriculum syllabus",
    }
    expanded = query
    for key, value in expansions.items():
        expanded = expanded.replace(key, value)
    return expanded